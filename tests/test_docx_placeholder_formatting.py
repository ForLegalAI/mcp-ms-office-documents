"""Tests for formatting preservation during template placeholder replacement.

``_replace_placeholder_in_paragraph`` rebuilds the paragraph around a placeholder.
These tests cover two guarantees:

* the formatting of the surrounding inline text (before/after the placeholder) is
  preserved instead of being flattened to plain runs, and
* the placeholder run's own bold/italic/underline/highlight (plus font/colour) is
  captured and applied to the replacement text, without clobbering formatting the
  markdown value asked for.
"""
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx_tools.dynamic_docx_tools import _replace_placeholder_in_paragraph  # noqa: E402


def _run_with_text(paragraph, text):
    """Return the first run whose text equals *text* (or None)."""
    return next((r for r in paragraph.runs if r.text == text), None)


# --- surrounding inline text keeps its formatting ---------------------------

def test_surrounding_runs_keep_their_formatting():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Important: ").bold = True
    p.add_run("{{val}}")
    p.add_run(" done").italic = True

    assert _replace_placeholder_in_paragraph(p, "{{val}}", "X", doc=doc)

    assert p.text == "Important: X done"
    assert _run_with_text(p, "Important: ").bold is True
    assert _run_with_text(p, " done").italic is True


def test_text_after_placeholder_with_color_is_preserved():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("{{val}}")
    after = p.add_run(" tail")
    after.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    assert _replace_placeholder_in_paragraph(p, "{{val}}", "head", doc=doc)

    assert p.text == "head tail"
    assert _run_with_text(p, " tail").font.color.rgb == RGBColor(0xFF, 0x00, 0x00)


# --- placeholder's own formatting is applied to the replacement -------------

def test_placeholder_bold_applied_to_replacement():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("{{val}}").bold = True

    assert _replace_placeholder_in_paragraph(p, "{{val}}", "hello", doc=doc)

    assert p.text == "hello"
    assert all(r.bold for r in p.runs)


def test_placeholder_italic_underline_highlight_applied():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("{{val}}")
    run.italic = True
    run.underline = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    assert _replace_placeholder_in_paragraph(p, "{{val}}", "hi", doc=doc)

    out = _run_with_text(p, "hi")
    assert out.italic is True
    assert out.underline is True
    assert out.font.highlight_color == WD_COLOR_INDEX.YELLOW


def test_placeholder_font_name_and_size_applied():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("{{val}}")
    run.font.name = "Courier New"
    run.font.size = Pt(14)

    assert _replace_placeholder_in_paragraph(p, "{{val}}", "hi", doc=doc)

    out = _run_with_text(p, "hi")
    assert out.font.name == "Courier New"
    assert out.font.size == Pt(14)


# --- markdown formatting in the value is not clobbered ----------------------

def test_value_markdown_combines_with_placeholder_format():
    # Placeholder is italic; the value asks for bold on part of it. The bold part
    # must stay bold (markdown) AND gain italic (placeholder) — neither wins out.
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("{{val}}").italic = True

    assert _replace_placeholder_in_paragraph(p, "{{val}}", "**strong**", doc=doc)

    out = _run_with_text(p, "strong")
    assert out.bold is True      # from the markdown value
    assert out.italic is True    # filled from the placeholder run


def test_placeholder_explicit_not_bold_is_propagated():
    # An explicit bold=False on the placeholder (e.g. to counteract a bold
    # paragraph style) should pass through to a plain replacement value.
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("{{val}}").bold = False

    assert _replace_placeholder_in_paragraph(p, "{{val}}", "hi", doc=doc)

    assert _run_with_text(p, "hi").bold is False


def test_value_markdown_bold_wins_over_placeholder_not_bold():
    # Even when the placeholder is explicitly not-bold, markdown bold in the value
    # must win (the value's formatting is never overridden).
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("{{val}}").bold = False

    assert _replace_placeholder_in_paragraph(p, "{{val}}", "**strong**", doc=doc)

    assert _run_with_text(p, "strong").bold is True


# --- placeholder split across runs (as Word often stores it) ----------------

def test_split_placeholder_preserves_bold_and_surrounding():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Hi ").bold = True
    r1 = p.add_run("{{va")
    r1.italic = True
    r2 = p.add_run("l}}")
    r2.italic = True
    p.add_run("!")

    assert _replace_placeholder_in_paragraph(p, "{{val}}", "there", doc=doc)

    assert p.text == "Hi there!"
    assert _run_with_text(p, "Hi ").bold is True
    assert _run_with_text(p, "there").italic is True   # from the (italic) placeholder run
