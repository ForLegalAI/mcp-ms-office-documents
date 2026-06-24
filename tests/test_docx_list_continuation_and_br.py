"""Tests for two related list-rendering fixes.

1. ``<br>`` bordering block content (a list/heading) is promoted to a real line
   break so the block is recognised — both in the base markdown pipeline and in
   the dynamic-template placeholder path. A prose ``<br>`` stays an inline soft
   break, and a ``<br>`` inside a table cell is untouched.

2. Ordered-list numbering continues across an intervening heading/blank line
   (e.g. procedural filings whose numbered paragraphs run 1, 2 under one heading
   and 3, 4 under the next). A running counter remembers where the previous list
   left off; a lone numbered line with no such predecessor still renders as prose
   (date disambiguation, see test_docx_ordered_list_date.py).

Numbering is asserted on the underlying XML because Word computes the visible
numbers at display time; they are not stored as paragraph text.
"""
import sys
from pathlib import Path

from docx import Document

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx_tools.markdown_processor import process_markdown_content  # noqa: E402
from docx_tools.patterns import expand_br_to_block_breaks  # noqa: E402
from docx_tools.dynamic_docx_tools import (  # noqa: E402
    _replace_placeholders_in_paragraph,
)


def _new_doc():
    default = project_root / "default_templates" / "default_docx_template.docx"
    assert default.exists(), f"Default template missing at {default}"
    return Document(str(default))


def _render(content):
    doc = _new_doc()
    start = len(doc.paragraphs)
    process_markdown_content(doc, content)
    return doc, doc.paragraphs[start:]


def _is_ordered(p):
    return bool(p.style.name and p.style.name.startswith("List Number"))


def _num_id_of(p):
    vals = p._p.xpath('.//w:numPr/w:numId/@w:val')
    return vals[0] if vals else None


def _start_override(doc, num_id):
    num = doc.part.numbering_part.element.num_having_numId(int(num_id))
    vals = num.xpath('./w:lvlOverride[@w:ilvl="0"]/w:startOverride/@w:val')
    return vals[0] if vals else None


# ---------------------------------------------------------------------------
# expand_br_to_block_breaks (unit)
# ---------------------------------------------------------------------------

def test_expand_br_splits_when_list_follows():
    assert expand_br_to_block_breaks("Intro:<br>1. A<br>2. B") == "Intro:\n1. A\n2. B"


def test_expand_br_splits_when_heading_follows():
    assert expand_br_to_block_breaks("Lead<br># Title") == "Lead\n# Title"


def test_expand_br_leaves_prose_soft_break_alone():
    # No block element after the <br>: keep it for the inline soft-break path.
    assert expand_br_to_block_breaks("Line A<br>Line B") == "Line A<br>Line B"


def test_expand_br_ignores_lone_number_that_is_a_date():
    # "23." after a <br> is not a genuine list (no continuation) -> not split.
    assert expand_br_to_block_breaks("On<br>23. brezna") == "On<br>23. brezna"


def test_expand_br_leaves_table_rows_untouched():
    table = "| a<br>b | c |"
    assert expand_br_to_block_breaks(table) == table


def test_expand_br_is_idempotent():
    once = expand_br_to_block_breaks("Intro:<br>1. A<br>2. B")
    assert expand_br_to_block_breaks(once) == once


# ---------------------------------------------------------------------------
# Fix #1 — <br> before a list (base pipeline)
# ---------------------------------------------------------------------------

def test_br_separated_list_becomes_numbered_items():
    doc, paras = _render("Vyrok soudu:<br>1. Prvni<br>2. Druhy<br>3. Treti")
    intro = [p for p in paras if p.text == "Vyrok soudu:"]
    items = [p for p in paras if _is_ordered(p)]
    assert len(intro) == 1 and not _is_ordered(intro[0])
    assert [p.text for p in items] == ["Prvni", "Druhy", "Treti"]


def test_prose_br_stays_single_soft_break_paragraph():
    doc, paras = _render("Line A<br>Line B")
    body = [p for p in paras if p.text.strip()]
    assert len(body) == 1
    assert "Line A" in body[0].text and "Line B" in body[0].text
    # A soft break (<w:br/>), not a new paragraph.
    assert body[0]._p.xpath('.//w:br')


def test_br_inside_table_cell_is_preserved():
    md = "| H1 | H2 |\n| --- | --- |\n| a<br>b | c |"
    doc, _ = _render(md)
    assert len(doc.tables) == 1
    cell = doc.tables[0].rows[1].cells[0]
    assert [p.text for p in cell.paragraphs] == ["a", "b"]


# ---------------------------------------------------------------------------
# Fix #2 — continuation across a heading
# ---------------------------------------------------------------------------

def test_numbering_continues_after_heading():
    md = (
        "## I. Uvod\n\n"
        "1. Prvni\n\n"
        "2. Druhy\n\n"
        "## II. Argumentace\n\n"
        "3. Treti\n\n"
        "4. Ctvrty\n"
    )
    doc, paras = _render(md)
    items = [p for p in paras if _is_ordered(p)]
    assert [p.text for p in items] == ["Prvni", "Druhy", "Treti", "Ctvrty"]

    first_run = _num_id_of(items[0])
    second_run = _num_id_of(items[2])
    assert _num_id_of(items[1]) == first_run
    assert _num_id_of(items[3]) == second_run
    assert first_run != second_run, "the continuation is a fresh numbering instance"
    # The continuation instance resumes the count at 3 via startOverride.
    assert _start_override(doc, first_run) == "1"
    assert _start_override(doc, second_run) == "3"


def test_single_continuation_item_after_heading_is_a_list():
    md = "1. Prvni\n\n2. Druhy\n\n## II\n\n3. Treti\n"
    doc, paras = _render(md)
    items = [p for p in paras if _is_ordered(p)]
    assert [p.text for p in items] == ["Prvni", "Druhy", "Treti"]
    assert _start_override(doc, _num_id_of(items[2])) == "3"


def test_new_list_starting_at_one_restarts_after_heading():
    md = "1. Prvni\n\n2. Druhy\n\n## II\n\n1. Novy\n\n2. Dalsi\n"
    doc, paras = _render(md)
    items = [p for p in paras if _is_ordered(p)]
    assert [p.text for p in items] == ["Prvni", "Druhy", "Novy", "Dalsi"]
    assert _start_override(doc, _num_id_of(items[2])) == "1"


def test_prose_between_lists_breaks_continuation():
    # An ordinary paragraph (not a heading) resets the run, so "3." with no
    # immediate sibling falls back to prose rather than joining the count.
    md = "1. Prvni\n\n2. Druhy\n\nNejaky odstavec.\n\n3. Treti\n"
    doc, paras = _render(md)
    assert any(p.text == "3. Treti" and not _is_ordered(p) for p in paras)


def test_standalone_non_one_pair_blank_separated_stays_prose():
    # No preceding list -> textually a date pair -> prose (date disambiguation).
    doc, paras = _render("5. Paty\n\n6. Sesty")
    assert all(not _is_ordered(p) for p in paras if p.text.strip())


def test_non_continuing_number_after_heading_stays_prose():
    # Continuation must match the running count exactly: after a list ends at 2
    # (run resumes at 3), a "23." under the next heading does not continue and so
    # — being blank-separated and not locally genuine — renders as prose.
    md = "1. Prvni\n\n2. Druhy\n\n## II\n\n23. brezna 2026\n"
    doc, paras = _render(md)
    assert any(p.text == "23. brezna 2026" and not _is_ordered(p) for p in paras)


# ---------------------------------------------------------------------------
# Fix #1b — <br> list inside a template placeholder
# ---------------------------------------------------------------------------

def _render_placeholder(value):
    doc = _new_doc()
    p = doc.add_paragraph()
    p.add_run("{{body}}")
    _replace_placeholders_in_paragraph(p, {"body": value}, doc=doc)
    return doc


def test_placeholder_br_separated_list_becomes_numbered_items():
    doc = _render_placeholder("Vyrok:<br>1. Prvni<br>2. Druhy")
    items = [p for p in doc.paragraphs if _is_ordered(p)]
    assert [p.text for p in items] == ["Prvni", "Druhy"]
    assert any(p.text == "Vyrok:" and not _is_ordered(p) for p in doc.paragraphs)


def test_placeholder_real_newline_list_still_works():
    doc = _render_placeholder("Intro\n1. Prvni\n2. Druhy")
    items = [p for p in doc.paragraphs if _is_ordered(p)]
    assert [p.text for p in items] == ["Prvni", "Druhy"]
