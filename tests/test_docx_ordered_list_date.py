"""Tests for ordered-list vs. date disambiguation.

A standalone numbered line such as a date ("23. června 2026") used to be misread
as an ordered-list item. ``ordered_list_is_genuine`` now requires a numbered line
to either start at 1 or have a continuation before it is treated as a list, so
dates on days 2–31 render as prose automatically. A day-1 date ("1. června 2026")
is indistinguishable from a one-item list and must be escaped ("1\\. ...").
"""
import sys
from pathlib import Path

from docx import Document

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from docx_tools.markdown_processor import process_markdown_content  # noqa: E402
from docx_tools.patterns import ordered_list_is_genuine  # noqa: E402


def _new_doc_with_default_styles():
    """Document from the project's default template (carries ``List Number``)."""
    default = project_root / "default_templates" / "default_docx_template.docx"
    assert default.exists(), f"Default template missing at {default}"
    return Document(str(default))


def _render(content):
    """Render *content* into a fresh document and return the new paragraphs."""
    doc = _new_doc_with_default_styles()
    start = len(doc.paragraphs)
    process_markdown_content(doc, content)
    return doc.paragraphs[start:]


def _is_ordered(paragraph):
    return bool(paragraph.style.name and paragraph.style.name.startswith("List Number"))


# --- helper-level tests -----------------------------------------------------

def test_lone_date_is_not_a_list():
    assert ordered_list_is_genuine(["23. června 2026"], 0) is False


def test_lone_non_one_number_is_not_a_list():
    assert ordered_list_is_genuine(["5. only item"], 0) is False


def test_list_starting_at_one_is_genuine():
    assert ordered_list_is_genuine(["1. only item"], 0) is True


def test_non_one_start_with_sibling_is_genuine():
    assert ordered_list_is_genuine(["5. Fifth", "6. Sixth"], 0) is True


def test_non_one_start_with_nested_child_is_genuine():
    assert ordered_list_is_genuine(["2. parent", "    - child"], 0) is True


def test_blank_line_ends_run_before_continuation():
    assert ordered_list_is_genuine(["5. Fifth", "", "6. Sixth"], 0) is False


def test_two_adjacent_dates_are_a_known_limitation():
    # Two consecutive date-like lines look exactly like a 2-item list; documented
    # limitation — escape them if they must stay prose.
    assert ordered_list_is_genuine(["23. června 2026", "24. července 2026"], 0) is True


# --- pipeline-level tests ---------------------------------------------------

def test_lone_date_renders_as_paragraph():
    paras = _render("23. června 2026")
    assert len(paras) == 1
    assert not _is_ordered(paras[0])
    assert paras[0].text == "23. června 2026"


def test_list_starting_at_five_still_renders_as_list():
    paras = _render("5. Fifth\n6. Sixth")
    assert all(_is_ordered(p) for p in paras)
    assert [p.text for p in paras] == ["Fifth", "Sixth"]


def test_single_item_list_starting_at_one_renders_as_list():
    paras = _render("1. Only item")
    assert len(paras) == 1
    assert _is_ordered(paras[0])
    assert paras[0].text == "Only item"


def test_two_item_list_renders_as_list():
    paras = _render("1. First\n2. Second")
    assert all(_is_ordered(p) for p in paras)
    assert [p.text for p in paras] == ["First", "Second"]


def test_escaped_day_one_date_renders_as_paragraph():
    paras = _render("1\\. června 2026")
    assert len(paras) == 1
    assert not _is_ordered(paras[0])
    assert paras[0].text == "1. června 2026"
